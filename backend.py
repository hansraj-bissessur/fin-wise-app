import uvicorn
import logging
from fastapi import FastAPI, File, UploadFile, Form, HTTPException, Depends
from pydantic import BaseModel
from typing import List, Optional
import io
import datetime

# Document Parsers
import pypdf
import docx
import openpyxl

# LangChain Components
from langchain_ollama import OllamaLLM as Ollama 
from langchain_ollama import OllamaEmbeddings 
from langchain_redis import RedisVectorStore as Redis 
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain.docstore.document import Document as LangChainDocument

# Redis client for admin tasks
import redis

# --- Warning Suppression ---
import warnings
from urllib3.exceptions import NotOpenSSLWarning


# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Initialize FastAPI app
app = FastAPI(
    title="Financial Literacy Chatbot API",
    description="API for RAG chatbot and document management."
)

# --- Data Models ---

class ChatRequest(BaseModel):
    message: str
    user_id: str = "anonymous"

class ChatResponse(BaseModel):
    response: str
    user_id: str
    confidence_score: float
    suggest_ticket: bool

class HealthResponse(BaseModel):
    status: str
    service: str
    timestamp: str

class UploadResponse(BaseModel):
    success: bool
    message: str
    documents_processed: int
    total_chunks_created: int

class ClearResponse(BaseModel):
    success: bool
    message: str

try:
    # Initialize Ollama embeddings
    embeddings = OllamaEmbeddings(model="nomic-embed-text")
    
    # Initialize Ollama chat model
    chat_model = Ollama(model="phi3:mini")

    # Initialize Redis vector store
    REDIS_URL = "redis://localhost:6379"
    INDEX_NAME = "financial_literacy_docs"
    
   
    vectorstore = Redis(
        redis_url=REDIS_URL,
        index_name=INDEX_NAME,
        embeddings=embeddings 
    )
    
    # Direct Redis client for admin tasks
    redis_client = redis.from_url(REDIS_URL)

    # Text Splitter (matches Java config)
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200
    )

    # System Prompt (matches Java config)
    SYSTEM_PROMPT = """
        You are a financial assistant. Provide concise, helpful answers using the context below.
        Keep responses under 150 words for mobile users.
        
        CONTEXT:
        {documents}
        """

except Exception as e:
    logger.error(f"Failed to initialize global components: {e}", exc_info=True)
    # In a real app, you might want to prevent startup
    vectorstore = None
    chat_model = None
    redis_client = None

# --- Dependency Check ---

def get_services():
    if not vectorstore or not chat_model or not redis_client:
        raise HTTPException(
            status_code=503, 
            detail="Core services (Ollama/Redis) are not available."
        )
    return vectorstore, chat_model, redis_client

# --- Document Processing Logic (from DocumentProcessingService.java) ---

def parse_pdf(contents: bytes) -> str:
    logger.info("Parsing PDF content")
    reader = pypdf.PdfReader(io.BytesIO(contents))
    text = ""
    for page in reader.pages:
        text += page.extract_text() or ""
    return text

def parse_docx(contents: bytes) -> str:
    logger.info("Parsing DOCX content")
    doc = docx.Document(io.BytesIO(contents))
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text

def parse_xlsx(contents: bytes) -> str:
    logger.info("Parsing XLSX content")
    wb = openpyxl.load_workbook(io.BytesIO(contents))
    text = ""
    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        text += f"Sheet: {sheet_name}\n"
        for row in sheet.iter_rows():
            row_values = [str(cell.value) for cell in row if cell.value is not None]
            if row_values:
                text += " | ".join(row_values) + "\n"
        text += "\n"
    return text

# --- FastAPI Endpoints (from ChatbotController.java) ---

@app.post("/api/v1/chatbot/chat", response_model=ChatResponse)
async def chat(request: ChatRequest, services = Depends(get_services)):
    """
    Main chatbot endpoint. Processes a user query using RAG.
    """
    vectorstore, chat_model, _ = services
    logger.info(f"Processing chat query for user: {request.user_id}")

    try:
        # Step 1: Retrieve relevant documents
        retriever = vectorstore.as_retriever(search_kwargs={"k": 3})
        relevant_docs = retriever.invoke(request.message)

        # Step 2: Prepare context
        doc_texts = [doc.page_content for doc in relevant_docs]
        documents_context = "\n---\n".join(doc_texts)
        if not documents_context:
            documents_context = "No specific context available."
            
        # Step 3: Prepare prompt and chain
        prompt_template = ChatPromptTemplate.from_messages([
            ("system", SYSTEM_PROMPT),
            ("human", "{query}")
        ])
        
        chain = prompt_template | chat_model | StrOutputParser()
        
        # Step 4: Invoke chain
        response_text = chain.invoke({
            "documents": documents_context,
            "query": request.message
        })
        
        # Step 5: Confidence scoring and ticket suggestion
        confidence_score = 0.2 if not relevant_docs else min(1.0, 0.5 + (len(relevant_docs) / 3.0 * 0.5))
        suggest_ticket = confidence_score < 0.5

        if suggest_ticket:
            response_text += "\n\n💡 **Need more help?** Contact our customer service team for personalized assistance."

        return ChatResponse(
            response=response_text,
            user_id=request.user_id,
            confidence_score=confidence_score,
            suggest_ticket=suggest_ticket
        )
    
    except Exception as e:
        logger.error(f"Error processing chat: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail="Internal server error occurred.")


@app.post("/api/v1/chatbot/documents/upload-multiple", response_model=UploadResponse)
async def upload_multiple_documents(
    files: List[UploadFile] = File(...),
    user_id: str = Form(...),
    services = Depends(get_services)
):
    """
    Uploads multiple documents, processes them, and adds them to the vector store.
    """
    vectorstore, _, _ = services
    logger.info(f"Processing {len(files)} files for user: {user_id}")

    if not files:
        raise HTTPException(status_code=400, detail="No files provided")

    total_chunks = 0
    docs_processed = 0
    all_langchain_docs = []
    
    parsers = {
        "application/pdf": parse_pdf,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": parse_docx, # .docx
        "application/msword": parse_docx, # .doc
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": parse_xlsx, # .xlsx
        "application/vnd.ms-excel": parse_xlsx # .xls
    }

    for file in files:
        if file.content_type not in parsers:
            logger.warning(f"Skipping unsupported file type: {file.filename} ({file.content_type})")
            continue
        
        try:
            contents = await file.read()
            parser = parsers[file.content_type]
            
            # 1. Extract text
            extracted_text = parser(contents)
            if not extracted_text or extracted_text.isspace():
                logger.warning(f"No text extracted from file: {file.filename}")
                continue

            # 2. Split text into chunks
            chunks = text_splitter.split_text(extracted_text)
            
            # 3. Create LangChain Documents with metadata
            for i, chunk in enumerate(chunks):
                metadata = {
                    "fileName": file.filename,
                    "fileType": file.content_type,
                    "userId": user_id,
                    "chunkIndex": i,
                    "totalChunks": len(chunks),
                    "category": "financial_literacy",
                    "uploadTimestamp": datetime.datetime.now().isoformat()
                }
                all_langchain_docs.append(LangChainDocument(page_content=chunk, metadata=metadata))
            
            total_chunks += len(chunks)
            docs_processed += 1
            
        except Exception as e:
            logger.error(f"Failed to process file {file.filename}: {e}", exc_info=True)
            raise HTTPException(status_code=500, detail=f"Error processing file: {file.filename}")

    # 4. Add all documents to vector store
    if all_langchain_docs:
        try:
            vectorstore.add_documents(all_langchain_docs)
            logger.info(f"Successfully added {total_chunks} chunks from {docs_processed} documents to Redis.")
        except Exception as e:
            logger.error(f"Failed to add documents to vector store: {e}", exc_info=True)
            raise HTTPException(status_code=500, detail="Error adding documents to vector store.")

    return UploadResponse(
        success=True,
        message=f"Successfully processed {docs_processed} files into {total_chunks} chunks.",
        documents_processed=docs_processed,
        total_chunks_created=total_chunks
    )


@app.delete("/api/v1/chatbot/documents/clear-all", response_model=ClearResponse)
async def clear_all_documents(admin_key: str, services = Depends(get_services)):
    """
    Clears all documents from the vector store.
    Matches the 'admin123' key from ChatbotController.java.
    """
    _, _, redis_client = services
    
    if admin_key != "admin123":
        raise HTTPException(status_code=403, detail="Unauthorized")
    
    try:
        redis_client.flushdb()
        logger.info("Successfully cleared all documents from Redis.")
        
        # Re-initialize the index schema after clearing
        vectorstore.schema = Redis.check_index_schema(redis_url=REDIS_URL, index_name=INDEX_NAME, embeddings=embeddings)
        
        return ClearResponse(
            success=True,
            message="All documents cleared successfully. The index is ready for new uploads."
        )
    except Exception as e:
        logger.error(f"Error clearing documents: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail="Failed to clear documents.")


@app.get("/health", response_model=HealthResponse)
async def health_check():
    """
    Provides a simple health check endpoint.
    """
    return HealthResponse(
        status="healthy",
        service="Financial Literacy Chatbot",
        timestamp=datetime.datetime.now().isoformat()
    )

if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)

